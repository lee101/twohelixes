"""Freeform documents expose real tables or an honest document frame."""

from __future__ import annotations

import base64
import json
from pathlib import Path
from typing import Any

import pytest

from twohelixes import auth, router, store
from twohelixes.ingest import MarkItDownRequired, ingest_path
from twohelixes.ingest import document as document_ingest
from twohelixes.routes import connectors


def _pdf(path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Table, TableStyle

    styles = getSampleStyleSheet()
    grid = TableStyle(
        [
            ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ]
    )
    first = Table(
        [["Region", "Revenue"], ["North", "10"], ["South", "20"]],
        colWidths=[180, 100],
    )
    first.setStyle(grid)
    second = Table(
        [["Product", "Stock"], ["Widget", "4"], ["Gadget", "8"]],
        colWidths=[180, 100],
    )
    second.setStyle(grid)
    SimpleDocTemplate(str(path), pagesize=letter).build(
        [
            Paragraph("Regional Sales", styles["Heading1"]),
            first,
            PageBreak(),
            Paragraph("Inventory", styles["Heading1"]),
            second,
        ]
    )


def _docx(path: Path) -> None:
    from docx import Document

    document = Document()
    document.add_heading("Customer Results", level=1)
    document.add_paragraph("This prose remains preserved in the original document.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Customer"
    table.rows[0].cells[1].text = "Score"
    for customer, score in (("Ada", "91"), ("Lin", "87")):
        cells = table.add_row().cells
        cells[0].text = customer
        cells[1].text = score
    document.save(path)


def test_pdf_with_two_tables_becomes_two_sheet_candidates(tmp_path: Path) -> None:
    path = tmp_path / "report.pdf"
    _pdf(path)

    result = ingest_path(path, allow_agent=False)

    assert len(result.report.sheets) == 2
    assert {sheet["name"] for sheet in result.report.sheets} == {
        "Regional Sales",
        "Inventory",
    }
    assert set(result.frame.columns) in (
        {"region", "revenue"},
        {"product", "stock"},
    )


def test_docx_table_and_prose_uses_the_heading_for_the_table(tmp_path: Path) -> None:
    path = tmp_path / "customers.docx"
    _docx(path)

    result = ingest_path(path, allow_agent=False)

    assert result.report.sheet == "Customer Results"
    assert result.frame.to_dict("records") == [
        {"customer": "Ada", "score": 91},
        {"customer": "Lin", "score": 87},
    ]
    assert any("converted the document" in note.casefold() for note in result.report.notes)


def test_html_table_is_extracted(tmp_path: Path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        "<html><body><h2>Quarterly totals</h2><p>Context.</p>"
        "<table><tr><th>Quarter</th><th>Total</th></tr>"
        "<tr><td>Q1</td><td>12</td></tr>"
        "<tr><td>Q2</td><td>18</td></tr></table></body></html>"
    )

    result = ingest_path(path, allow_agent=False)

    assert result.report.sheet == "Quarterly totals"
    assert result.frame["total"].tolist() == [12, 18]


def test_markdown_table_is_extracted(tmp_path: Path) -> None:
    path = tmp_path / "notes.md"
    path.write_text(
        "# Release metrics\n\n"
        "| Version | Installs |\n"
        "| --- | ---: |\n"
        "| 1.0 | 120 |\n"
        "| 1.1 | 180 |\n"
    )

    result = ingest_path(path, allow_agent=False)

    assert result.report.sheet == "Release metrics"
    assert result.frame["installs"].tolist() == [120, 180]


def test_document_without_tables_becomes_document_dataset(tmp_path: Path) -> None:
    path = tmp_path / "brief.md"
    path.write_text(
        "# Overview\n\nThis is the first paragraph.\n\n"
        "This is the second paragraph with more words.\n\n"
        "## Findings\n\nRevenue improved."
    )

    result = ingest_path(path, allow_agent=False)

    assert result.report.dataset_kind == "document"
    assert list(result.frame.columns) == [
        "heading",
        "text",
        "section_index",
        "word_count",
    ]
    assert len(result.frame) == 3
    assert result.frame["heading"].tolist() == ["Overview", "Overview", "Findings"]
    assert any("No tables were found" in note for note in result.report.notes)


def test_rtf_control_words_do_not_leak_into_document_rows(tmp_path: Path) -> None:
    path = tmp_path / "brief.rtf"
    path.write_text(
        r"{\rtf1\ansi\b Overview\b0\par This is readable prose.\par}",
        encoding="ascii",
    )

    result = ingest_path(path, allow_agent=False)

    assert result.report.dataset_kind == "document"
    assert "readable prose" in " ".join(result.frame["text"]).casefold()
    assert "\\rtf" not in " ".join(result.frame["text"])


def test_markitdown_absent_degrades_cleanly(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    path = tmp_path / "brief.pdf"
    path.write_bytes(b"%PDF-not-read")
    real_find_spec = document_ingest.importlib.util.find_spec

    def missing(name: str, *args: Any, **kwargs: Any) -> Any:
        if name == "markitdown":
            return None
        return real_find_spec(name, *args, **kwargs)

    monkeypatch.setattr(document_ingest.importlib.util, "find_spec", missing)

    with pytest.raises(MarkItDownRequired, match="needs markitdown"):
        ingest_path(path, allow_agent=False)


@pytest.fixture
def upload_identity(tmp_path: Path, monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setenv("TWOHELIXES_DATA_DIR", str(tmp_path))
    store.close()
    store._initialised = False
    store.init()
    user = store.create_user("document-upload@example.com")
    yield auth.Identity(user_id=user["id"], email=user["email"], plan="free")
    store.close()
    store._initialised = False


def _ctx(identity: auth.Identity, body: dict[str, Any]) -> router.Context:
    context = router.build_context("POST", "/v1/upload", "", json.dumps(body), "")
    context.user = identity
    return context


def test_oversized_upload_is_rejected_before_conversion(
    upload_identity: auth.Identity, monkeypatch: pytest.MonkeyPatch
) -> None:
    monkeypatch.setattr(connectors, "MAX_UPLOAD_BYTES", 8)
    response = connectors.upload(
        _ctx(
            upload_identity,
            {
                "filename": "large.pdf",
                "content_base64": base64.b64encode(b"012345678").decode(),
            },
        )
    )

    assert response.status == 413
    assert response.body["error"] == "file_too_large"


def test_unsupported_upload_lists_the_supported_families(
    upload_identity: auth.Identity,
) -> None:
    response = connectors.upload(
        _ctx(
            upload_identity,
            {
                "filename": "archive.unknown",
                "content_base64": base64.b64encode(b"data").decode(),
            },
        )
    )

    assert response.status == 415
    assert response.body["error"] == "unsupported_type"
    assert "PDF" in response.body["detail"]
    assert "fixed-width" in response.body["detail"]


def test_pasted_google_sheet_link_uses_the_normal_upload_path(
    upload_identity: auth.Identity, monkeypatch: pytest.MonkeyPatch
) -> None:
    seen: dict[str, str] = {}

    def fetch(url: str, **kwargs: Any) -> tuple[bytes, str]:
        seen["url"] = url
        return b"city,sales\nA,1\nB,2\n", "google-sheet.csv"

    monkeypatch.setattr(connectors, "fetch_google_sheet", fetch)
    response = connectors.upload(
        _ctx(
            upload_identity,
            {
                "url": (
                    "https://docs.google.com/spreadsheets/d/sheet-id/"
                    "edit?usp=sharing"
                )
            },
        )
    )

    assert response.status == 201
    assert response.body["rows"] == 2
    assert response.body["columns"] == ["city", "sales"]
    assert seen["url"].startswith("https://docs.google.com/spreadsheets/")
    row = store.one(
        "SELECT storage, raw_storage, shape_report FROM datasets WHERE id = ?",
        (response.body["dataset_id"],),
    )
    assert Path(row["storage"]).suffix == ".parquet"
    assert Path(row["raw_storage"]).suffix == ".csv"
